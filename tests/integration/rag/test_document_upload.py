"""
Integration Tests for Document Upload with Multi-Format Support.

Tests the file parsing and document ingestion for various file formats (PDF, DOCX, Excel, HTML, Markdown, Text).

Setup:
    - Requires OpenAI API key for embeddings
    - Requires running Qdrant instance (http://localhost:6333)
    - Tests file parsing without full ingestion to avoid API costs

Run:
    pytest tests/integration/rag/test_document_upload.py -v
"""

import pytest
import uuid
import os
import io
from pathlib import Path

# Skip if not in integration test mode
pytestmark = pytest.mark.integration


# =============================================================================
# Test 1: Text File Parsing
# =============================================================================


@pytest.mark.asyncio
async def test_parse_text_file():
    """Test plain text file parsing."""
    from cortex.rag.parsers import parse_file

    content = "This is a sample text document for testing.\n\nIt has multiple lines."
    content_bytes = content.encode("utf-8")

    result = await parse_file(content_bytes, "test.txt")

    assert result == content
    assert "sample text document" in result


# =============================================================================
# Test 2: Markdown File Parsing
# =============================================================================


@pytest.mark.asyncio
async def test_parse_markdown_file():
    """Test markdown file parsing."""
    from cortex.rag.parsers import parse_file

    content = """# Header 1

## Header 2

This is **bold** and *italic* text.

- List item 1
- List item 2

[Link](https://example.com)
"""
    content_bytes = content.encode("utf-8")

    result = await parse_file(content_bytes, "test.md")

    # Should preserve markdown syntax (not convert to HTML in default mode)
    assert "# Header 1" in result or "Header 1" in result
    assert "bold" in result
    assert "List item 1" in result


# =============================================================================
# Test 3: HTML File Parsing
# =============================================================================


@pytest.mark.asyncio
async def test_parse_html_file():
    """Test HTML file parsing (strips tags, extracts text)."""
    from cortex.rag.parsers import parse_file

    html_content = """<!DOCTYPE html>
<html>
<head>
    <title>Test Page</title>
    <script>console.log('should be removed');</script>
    <style>body { color: red; }</style>
</head>
<body>
    <h1>Main Title</h1>
    <p>Paragraph with <strong>bold text</strong>.</p>
    <ul>
        <li>Item 1</li>
        <li>Item 2</li>
    </ul>
</body>
</html>"""
    content_bytes = html_content.encode("utf-8")

    result = await parse_file(content_bytes, "test.html")

    # Should extract text, remove HTML tags
    assert "Main Title" in result
    assert "bold text" in result
    assert "Item 1" in result

    # Should NOT contain HTML tags or scripts
    assert "<html>" not in result
    assert "<script>" not in result
    assert "console.log" not in result


# =============================================================================
# Test 4: PDF File Parsing
# =============================================================================


@pytest.mark.asyncio
async def test_parse_pdf_file():
    """Test PDF file parsing."""
    from cortex.rag.parsers import parse_file, PDF_AVAILABLE

    if not PDF_AVAILABLE:
        pytest.skip("pypdf not installed")

    # Create a minimal PDF with text
    try:
        from pypdf import PdfWriter
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
    except ImportError:
        pytest.skip("reportlab not installed (needed to create test PDF)")

    # Create PDF in memory
    pdf_buffer = io.BytesIO()
    c = canvas.Canvas(pdf_buffer, pagesize=letter)
    c.drawString(100, 750, "This is a test PDF document.")
    c.drawString(100, 700, "It contains sample text for parsing.")
    c.showPage()
    c.save()

    pdf_bytes = pdf_buffer.getvalue()

    result = await parse_file(pdf_bytes, "test.pdf")

    # Should extract text from PDF
    assert "test PDF document" in result or "PDF" in result


# =============================================================================
# Test 5: DOCX File Parsing
# =============================================================================


@pytest.mark.asyncio
async def test_parse_docx_file():
    """Test DOCX file parsing."""
    from cortex.rag.parsers import parse_file, DOCX_AVAILABLE

    if not DOCX_AVAILABLE:
        pytest.skip("python-docx not installed")

    # Create a minimal DOCX with text
    from docx import Document

    doc = Document()
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("This is a sample paragraph in a DOCX file.")
    doc.add_paragraph("It has multiple paragraphs.")

    # Save to bytes
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_bytes = docx_buffer.getvalue()

    result = await parse_file(docx_bytes, "test.docx")

    # Should extract text from DOCX
    assert "Test Document" in result
    assert "sample paragraph" in result
    assert "multiple paragraphs" in result


# =============================================================================
# Test 6: Excel File Parsing
# =============================================================================


@pytest.mark.asyncio
async def test_parse_excel_file():
    """Test Excel file parsing."""
    from cortex.rag.parsers import parse_file, EXCEL_AVAILABLE

    if not EXCEL_AVAILABLE:
        pytest.skip("openpyxl not installed")

    # Create a minimal Excel file with data
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "TestSheet"

    # Add header and data
    ws.append(["Name", "Age", "City"])
    ws.append(["Alice", 30, "New York"])
    ws.append(["Bob", 25, "San Francisco"])
    ws.append(["Charlie", 35, "Seattle"])

    # Save to bytes
    excel_buffer = io.BytesIO()
    wb.save(excel_buffer)
    excel_bytes = excel_buffer.getvalue()

    result = await parse_file(excel_bytes, "test.xlsx")

    # Should extract cell data
    assert "Name" in result
    assert "Age" in result
    assert "Alice" in result
    assert "New York" in result
    assert "Bob" in result


# =============================================================================
# Test 7: Unsupported File Format
# =============================================================================


@pytest.mark.asyncio
async def test_unsupported_file_format():
    """Test that unsupported file formats raise UnsupportedFileTypeError."""
    from cortex.rag.parsers import parse_file, UnsupportedFileTypeError

    # PNG magic bytes
    png_bytes = b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR"

    with pytest.raises(UnsupportedFileTypeError) as exc_info:
        await parse_file(png_bytes, "test.png")

    assert "Unsupported file type" in str(exc_info.value)
    assert ".png" in str(exc_info.value)


# =============================================================================
# Test 8: File Parsing Error (Empty PDF)
# =============================================================================


@pytest.mark.asyncio
async def test_file_parsing_error():
    """Test that malformed files raise FileParsingError."""
    from cortex.rag.parsers import parse_file, FileParsingError, PDF_AVAILABLE

    if not PDF_AVAILABLE:
        pytest.skip("pypdf not installed")

    # Invalid PDF bytes (just PDF header without body)
    invalid_pdf = b"%PDF-1.4\n"

    with pytest.raises(FileParsingError) as exc_info:
        await parse_file(invalid_pdf, "test.pdf")

    assert "Failed to parse PDF" in str(exc_info.value)


# =============================================================================
# Test 9: Format Detection by Extension
# =============================================================================


@pytest.mark.asyncio
async def test_format_detection_by_extension():
    """Test that format is detected correctly by file extension."""
    from cortex.rag.parsers import parse_file

    # Same content, different extensions
    content_bytes = b"Sample text content"

    # .txt should work
    result_txt = await parse_file(content_bytes, "file.txt")
    assert result_txt == "Sample text content"

    # .log should work (treated as text)
    result_log = await parse_file(content_bytes, "file.log")
    assert result_log == "Sample text content"

    # .text should work
    result_text = await parse_file(content_bytes, "file.text")
    assert result_text == "Sample text content"


# =============================================================================
# Test 10: Non-ASCII Characters (UTF-8)
# =============================================================================


@pytest.mark.asyncio
async def test_parse_file_with_unicode():
    """Test parsing files with non-ASCII characters."""
    from cortex.rag.parsers import parse_file

    # Content with emojis and non-ASCII
    content = "Hello 世界! 🌍 Testing UTF-8 support: café, résumé, naïve"
    content_bytes = content.encode("utf-8")

    result = await parse_file(content_bytes, "test.txt")

    assert "世界" in result
    assert "🌍" in result
    assert "café" in result
    assert "résumé" in result


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
